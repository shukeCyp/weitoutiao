from __future__ import annotations

import json
import logging
import os
import re
import subprocess
import webbrowser
from datetime import datetime
from pathlib import Path

import webview
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from peewee import IntegrityError

from .db import (
    count_benchmark_accounts,
    count_monitor_runs,
    count_monitored_articles,
    count_settings,
    create_benchmark_account,
    create_monitor_run,
    delete_benchmark_account,
    delete_benchmark_accounts,
    delete_monitored_article,
    export_benchmark_accounts,
    finish_monitor_run,
    get_benchmark_account_by_id,
    get_benchmark_account_by_url,
    get_monitored_article_by_id,
    get_setting,
    import_benchmark_accounts,
    list_article_ids_for_batch,
    list_benchmark_account_options,
    list_benchmark_accounts,
    list_completed_article_ids,
    list_monitored_articles,
    persist_monitored_articles,
    set_setting,
    soft_delete_all_monitored_articles,
    update_benchmark_account,
)
from .article_rewrite_pipeline import ArticleRewritePipeline
from .monitoring import AccountMonitor, MonitoringThresholds, run_batch_account_monitoring
from .paths import DB_DIR, LOG_DIR, SETTINGS_DB_FILE
from .title_rewriter import ArticleRewriter, TitleRewriter

REWRITE_WORKER_COUNT_KEY = "rewrite.worker_count"
IMAGE_GEN_WORKER_COUNT_KEY = "image_gen.worker_count"
DEFAULT_REWRITE_WORKER_COUNT = 1
DEFAULT_IMAGE_GEN_WORKER_COUNT = 3
MAX_WORKER_COUNT = 16

logger = logging.getLogger("weitoutiao.api")


def _open_folder(path: Path) -> bool:
    path.mkdir(parents=True, exist_ok=True)
    try:
        subprocess.Popen(["open", str(path)], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return True
    except Exception:
        return False


class Api:
    """API surface for frontend/backend integration."""

    def __init__(self) -> None:
        self.window: webview.Window | None = None

    def attach_window(self, window: webview.Window) -> None:
        self.window = window

    def ping(self) -> str:
        return "pong"

    def download_rewritten_article(self, content: str) -> dict[str, str]:
        normalized_content = content.strip()
        if not normalized_content:
            return {"status": "error", "message": "没有可下载的改写内容。"}
        if self.window is None:
            return {"status": "error", "message": "下载窗口尚未初始化。"}

        logger.info("开始下载改写文章")
        try:
            result = self.window.create_file_dialog(
                webview.SAVE_DIALOG,
                save_filename="改写文章.txt",
                file_types=("文本文件 (*.txt)", "全部文件 (*.*)"),
            )
            if not result:
                logger.info("下载改写文章已取消")
                return {"status": "cancel"}

            selected_path = Path(result if isinstance(result, str) else result[0])
            selected_path.write_text(normalized_content, encoding="utf-8")
            logger.info("下载改写文章完成 path=%s", selected_path)
            return {"status": "success", "path": str(selected_path)}
        except Exception as exc:
            logger.exception("下载改写文章失败")
            return {"status": "error", "message": str(exc) or "下载改写文章失败。"}

    def download_article_docx(self, article_id: int) -> dict:
        """生成 Word 文档并弹出保存对话框。图片按约定位置插入段落间。"""
        if self.window is None:
            return {"status": "error", "message": "下载窗口尚未初始化。"}

        article = get_monitored_article_by_id(article_id)
        if article is None:
            return {"status": "error", "message": "未找到对应文章。"}

        title = (article.rewritten_title or "").strip()
        intro = (article.rewritten_intro or "").strip()
        sections_raw = article.rewritten_article
        paths_raw = article.image_paths

        if not title or not sections_raw:
            return {"status": "error", "message": "文章改写内容不完整，无法下载。"}

        try:
            sections: list[str] = json.loads(sections_raw)
        except (json.JSONDecodeError, ValueError):
            return {"status": "error", "message": "改写正文格式错误，无法下载。"}

        try:
            image_paths: list[str | None] = json.loads(paths_raw) if paths_raw else []
        except (json.JSONDecodeError, ValueError):
            image_paths = []

        safe_name = re.sub(r'[\\/:*?"<>|]', "_", title)[:80] or "改写文章"
        logger.info("开始下载 Word 文档 article_id=%s title=%r", article_id, title)

        try:
            result = self.window.create_file_dialog(
                webview.SAVE_DIALOG,
                save_filename=f"{safe_name}.docx",
                file_types=("Word 文档 (*.docx)", "全部文件 (*.*)"),
            )
            if not result:
                logger.info("下载 Word 文档已取消 article_id=%s", article_id)
                return {"status": "cancel"}

            save_path = Path(result if isinstance(result, str) else result[0])
            if not save_path.suffix:
                save_path = save_path.with_suffix(".docx")

            self._build_docx(title, intro, sections, image_paths, save_path)
            logger.info("Word 文档保存完成 path=%s", save_path)
            return {"status": "success", "path": str(save_path)}
        except Exception as exc:
            logger.exception("下载 Word 文档失败 article_id=%s", article_id)
            return {"status": "error", "message": str(exc) or "下载失败。"}

    @staticmethod
    def _build_docx(
        title: str,
        intro: str,
        sections: list[str],
        image_paths: list[str | None],
        save_path: Path,
    ) -> None:
        """
        构建 Word 文档。排版顺序：
          标题 → 导语 → 图片1 → 01+正文 → 02+正文 → … → 图片2（section 4 后）→ … → 图片3（section 7 后）
        """
        # 防御性去除 "########标题" 前缀（旧数据兼容）
        title_lines = [l for l in title.splitlines() if not l.strip().startswith("########")]
        title = " ".join(title_lines).strip() or title

        # 图片2 在 section index 4 后，图片3 在 section index 7 后；图片1 紧跟导语
        INSERT_AFTER: dict[int, int] = {4: 1, 7: 2}

        doc = Document()

        # ── 页面边距 ──────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # ── 标题 ──────────────────────────────────────────────
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x3D)

        doc.add_paragraph()  # 标题后留一行间距

        def _add_text_paras(text: str) -> None:
            """将文本按换行符拆分，每行作为独立 Word 段落写入，两端对齐 12pt。"""
            # 统一处理 \n\n 和 \n：每一行独立成段，避免段内软换行被 JUSTIFY 拉伸
            lines = [line.strip() for line in text.replace("\n\n", "\n").split("\n") if line.strip()]
            for line in lines:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)

        def _insert_image(img_list_idx: int) -> None:
            path_val = image_paths[img_list_idx] if img_list_idx < len(image_paths) else None
            if not path_val:
                return
            img_path = Path(path_val)
            # 兼容旧数据：扩展名不匹配时尝试其他常见格式
            if not img_path.exists():
                for alt_ext in (".jpeg", ".jpg", ".png", ".webp"):
                    alt = img_path.with_suffix(alt_ext)
                    if alt.exists():
                        img_path = alt
                        break
                else:
                    logger.warning("图片文件不存在，跳过插入 path=%s", img_path)
                    return
            try:
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(str(img_path), width=Inches(5.5))
                doc.add_paragraph()  # 图片后留空行
            except Exception:
                logger.warning("插入图片失败 path=%s", img_path, exc_info=True)

        # ── 导语 ──────────────────────────────────────────────
        if intro:
            _add_text_paras(intro)

        # ── 图片 1（导语之后，正文之前）─────────────────────
        _insert_image(0)

        # ── 正文段落：序号 + 内容 + 图片 ─────────────────────
        for section_idx, text in enumerate(sections):
            # 序号标签：01, 02, … 10
            label_num = f"{section_idx + 1:02d}"
            label_para = doc.add_paragraph(label_num)
            label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            label_para.paragraph_format.space_before = Pt(6)
            for run in label_para.runs:
                run.bold = True
                run.font.size = Pt(12)

            # 正文内容
            _add_text_paras(text)

            # 图片2、图片3 插入
            if section_idx in INSERT_AFTER:
                _insert_image(INSERT_AFTER[section_idx])

        save_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(save_path))

    def batch_download_articles_docx(self, article_ids: list[int]) -> dict:
        """批量生成 Word 文档，弹出文件夹选择对话框，所有文档保存到该目录。"""
        if self.window is None:
            return {"status": "error", "message": "下载窗口尚未初始化。"}
        if not article_ids:
            return {"status": "error", "message": "未选择任何文章。"}

        logger.info("开始批量下载 Word 文档 count=%s", len(article_ids))

        try:
            result = self.window.create_file_dialog(webview.FOLDER_DIALOG)
            if not result:
                logger.info("批量下载 Word 文档已取消")
                return {"status": "cancel"}

            folder = Path(result[0] if isinstance(result, (list, tuple)) else result)
            folder.mkdir(parents=True, exist_ok=True)

            succeeded = 0
            failed = 0
            skipped = 0
            errors: list[str] = []

            for article_id in article_ids:
                article = get_monitored_article_by_id(article_id)
                if article is None:
                    skipped += 1
                    continue

                title = (article.rewritten_title or "").strip()
                intro = (article.rewritten_intro or "").strip()
                sections_raw = article.rewritten_article
                paths_raw = article.image_paths

                if not title or not sections_raw:
                    skipped += 1
                    continue

                try:
                    sections: list[str] = json.loads(sections_raw)
                    image_paths: list[str | None] = json.loads(paths_raw) if paths_raw else []
                except (json.JSONDecodeError, ValueError):
                    skipped += 1
                    continue

                safe_name = re.sub(r'[\\/:*?"<>|]', "_", title)[:80] or f"文章_{article_id}"
                save_path = folder / f"{safe_name}.docx"

                # 如果同名文件已存在则加序号避免覆盖
                if save_path.exists():
                    stem = save_path.stem
                    suffix = save_path.suffix
                    counter = 1
                    while save_path.exists():
                        save_path = folder / f"{stem}_{counter}{suffix}"
                        counter += 1

                try:
                    self._build_docx(title, intro, sections, image_paths, save_path)
                    succeeded += 1
                    logger.info("批量下载：文档保存成功 path=%s", save_path)
                except Exception as exc:
                    failed += 1
                    errors.append(f"{title[:20]}: {exc}")
                    logger.warning("批量下载：文档保存失败 article_id=%s error=%s", article_id, exc)

            logger.info(
                "批量下载完成 succeeded=%s failed=%s skipped=%s folder=%s",
                succeeded, failed, skipped, folder,
            )
            return {
                "status": "success",
                "folder": str(folder),
                "succeeded": succeeded,
                "failed": failed,
                "skipped": skipped,
                "errors": errors[:5],
            }
        except Exception as exc:
            logger.exception("批量下载 Word 文档失败")
            return {"status": "error", "message": str(exc) or "批量下载失败。"}

    def get_setting(self, key: str) -> str | None:
        return get_setting(key)

    def set_setting(self, key: str, value: str) -> bool:
        logger.info("开始保存设置 key=%s", key)
        try:
            result = set_setting(key, value)
            logger.info("设置保存成功 key=%s", key)
            return result
        except Exception:
            logger.exception("设置保存失败 key=%s", key)
            raise

    def get_monitoring_status(self) -> dict[str, str | int]:
        LOG_DIR.mkdir(parents=True, exist_ok=True)
        DB_DIR.mkdir(parents=True, exist_ok=True)

        log_files = [p for p in LOG_DIR.glob("app.log*") if p.is_file()]
        log_size = sum(p.stat().st_size for p in log_files)
        db_size = SETTINGS_DB_FILE.stat().st_size if SETTINGS_DB_FILE.exists() else 0

        return {
            "logDir": str(LOG_DIR),
            "dbFile": str(SETTINGS_DB_FILE),
            "logFiles": len(log_files),
            "logSizeBytes": log_size,
            "dbSizeBytes": db_size,
            "settingCount": count_settings(),
            "benchmarkAccountCount": count_benchmark_accounts(),
            "monitorRunCount": count_monitor_runs(),
            "monitoredArticleCount": count_monitored_articles(),
        }

    def list_benchmark_accounts(self, page: int, page_size: int) -> dict:
        return list_benchmark_accounts(page, page_size)

    def list_benchmark_account_options(self) -> list[dict]:
        return list_benchmark_account_options()

    def create_benchmark_account(self, url: str) -> dict:
        normalized_url = url.strip()
        logger.info("开始创建对标账号 url=%s", normalized_url)
        try:
            result = create_benchmark_account(normalized_url)
            logger.info("对标账号创建成功 id=%s url=%s", result.get("id"), normalized_url)
            return result
        except IntegrityError as exc:
            logger.warning("对标账号已存在 url=%s", normalized_url)
            raise ValueError("该对标账号已存在。") from exc
        except Exception:
            logger.exception("对标账号创建失败 url=%s", normalized_url)
            raise

    def update_benchmark_account(self, account_id: int, url: str) -> dict:
        normalized_url = url.strip()
        logger.info("开始更新对标账号 id=%s url=%s", account_id, normalized_url)
        try:
            result = update_benchmark_account(account_id, normalized_url)
            logger.info("对标账号更新成功 id=%s url=%s", account_id, normalized_url)
            return result
        except IntegrityError as exc:
            logger.warning("对标账号链接冲突 id=%s url=%s", account_id, normalized_url)
            raise ValueError("该对标账号已存在。") from exc
        except Exception:
            logger.exception("对标账号更新失败 id=%s", account_id)
            raise

    def delete_benchmark_account(self, account_id: int) -> bool:
        logger.info("开始删除对标账号 id=%s", account_id)
        try:
            deleted = delete_benchmark_account(account_id)
            logger.info("对标账号删除完成 id=%s success=%s", account_id, deleted)
            return deleted
        except Exception:
            logger.exception("对标账号删除失败 id=%s", account_id)
            raise

    def delete_benchmark_accounts(self, account_ids: list[int]) -> int:
        logger.info("开始批量删除对标账号 ids=%s", account_ids)
        try:
            deleted_count = delete_benchmark_accounts(account_ids)
            logger.info("批量删除对标账号完成 count=%s", deleted_count)
            return deleted_count
        except Exception:
            logger.exception("批量删除对标账号失败 ids=%s", account_ids)
            raise

    def import_benchmark_accounts(self, payload: str) -> dict:
        logger.info("开始导入对标账号")
        try:
            result = import_benchmark_accounts(payload)
            logger.info(
                "导入对标账号完成 created=%s updated=%s skipped=%s total=%s",
                result.get("created"),
                result.get("updated"),
                result.get("skipped"),
                result.get("total"),
            )
            return result
        except Exception:
            logger.exception("导入对标账号失败")
            raise

    def export_benchmark_accounts(self) -> list[dict]:
        logger.info("开始导出对标账号")
        try:
            result = export_benchmark_accounts()
            logger.info("导出对标账号完成 count=%s", len(result))
            return result
        except Exception:
            logger.exception("导出对标账号失败")
            raise

    def run_account_monitor(self, payload: dict) -> dict:
        url = str(payload.get("url", "")).strip()
        start_time_raw = payload.get("startTime")
        end_time_raw = payload.get("endTime")
        single_capture = bool(payload.get("singleCapture", False))
        benchmark_account_id = payload.get("benchmarkAccountId")

        logger.info(
            "开始执行账号监控 url=%s start_time=%s end_time=%s single_capture=%s benchmark_account_id=%s",
            url,
            start_time_raw,
            end_time_raw,
            single_capture,
            benchmark_account_id,
        )
        monitor_run = None
        try:
            start_time = datetime.fromisoformat(start_time_raw) if start_time_raw else None
            end_time = datetime.fromisoformat(end_time_raw) if end_time_raw else None

            benchmark_account = None
            if benchmark_account_id is not None:
                benchmark_account = get_benchmark_account_by_id(int(benchmark_account_id))
            if benchmark_account is None and url:
                benchmark_account = get_benchmark_account_by_url(url)
            if benchmark_account is None and url:
                benchmark_account = get_benchmark_account_by_url(create_benchmark_account(url).get("url"))

            if benchmark_account is None:
                raise ValueError("未找到可关联的对标账号。")

            monitor_run = create_monitor_run(benchmark_account, url)

            monitor = AccountMonitor(
                url=url,
                start_time=start_time,
                end_time=end_time,
                single_capture=single_capture,
            )
            result = monitor.run()
            saved_count = persist_monitored_articles(benchmark_account, monitor_run, result.get("articles") or [])
            finish_monitor_run(
                monitor_run,
                status="warning" if result.get("warning") else "success",
                warning=result.get("warning"),
                article_count=saved_count,
            )
            result["savedCount"] = saved_count
            logger.info(
                "账号监控执行完成 url=%s capture_count=%s article_count=%s filtered_article_count=%s saved_count=%s warning=%s single_capture=%s",
                url,
                result.get("captureCount"),
                result.get("articleCount"),
                result.get("filteredArticleCount"),
                saved_count,
                result.get("warning"),
                single_capture,
            )
            return result
        except Exception as exc:
            if monitor_run is not None:
                finish_monitor_run(monitor_run, status="failed", warning=str(exc), article_count=0)
            logger.exception("账号监控执行失败 url=%s", url)
            raise

    def run_article_monitoring(self, payload: dict) -> dict:
        benchmark_account_ids = payload.get("benchmarkAccountIds") or []
        start_time_raw = payload.get("startTime")
        end_time_raw = payload.get("endTime")
        thresholds = MonitoringThresholds(
            min_play_count=self._parse_non_negative_int(payload.get("minPlayCount")),
            min_digg_count=self._parse_non_negative_int(payload.get("minDiggCount")),
            min_forward_count=self._parse_non_negative_int(payload.get("minForwardCount")),
        )

        logger.info(
            "开始执行批量文章监控 benchmark_account_ids=%s start_time=%s end_time=%s min_play_count=%s min_digg_count=%s min_forward_count=%s",
            benchmark_account_ids,
            start_time_raw,
            end_time_raw,
            thresholds.min_play_count,
            thresholds.min_digg_count,
            thresholds.min_forward_count,
        )

        try:
            start_time = datetime.fromisoformat(start_time_raw) if start_time_raw else None
            end_time = datetime.fromisoformat(end_time_raw) if end_time_raw else None
            result = run_batch_account_monitoring(
                benchmark_account_ids=[int(account_id) for account_id in benchmark_account_ids],
                start_time=start_time,
                end_time=end_time,
                thresholds=thresholds,
            )
            logger.info(
                "批量文章监控完成 requested=%s succeeded=%s failed=%s warning=%s",
                result.get("requestedCount"),
                result.get("succeededCount"),
                result.get("failedCount"),
                result.get("warningCount"),
            )
            return result
        except Exception:
            logger.exception("批量文章监控执行失败 benchmark_account_ids=%s", benchmark_account_ids)
            raise

    def _parse_non_negative_int(self, value: object) -> int:
        if value is None or value == "":
            return 0
        try:
            return max(0, int(str(value).strip()))
        except (TypeError, ValueError):
            raise ValueError("监控筛选条件必须是大于等于 0 的整数。")

    def rewrite_article(self, content: str, template_key: str) -> str:
        logger.info("开始改写文章 template_key=%s", template_key)
        try:
            result = ArticleRewriter().rewrite(content, template_key)
            logger.info("文章改写完成 template_key=%s length=%s", template_key, len(result))
            return result
        except Exception:
            logger.exception("文章改写失败 template_key=%s", template_key)
            raise

    def list_monitored_articles(self, filters: dict, page: int, page_size: int) -> dict:
        return list_monitored_articles(filters, page, page_size)

    def list_article_ids_for_batch(self, filters: dict, skip_rewritten: bool) -> list[int]:
        logger.info("查询批量改写 ID skip_rewritten=%s", skip_rewritten)
        return list_article_ids_for_batch(filters, skip_rewritten=skip_rewritten)

    def list_completed_article_ids(self, filters: dict) -> list[int]:
        logger.info("查询批量下载 ID（已完成改写）")
        return list_completed_article_ids(filters)

    def soft_delete_all_monitored_articles(self) -> int:
        logger.info("开始全量软删除文章")
        try:
            deleted_count = soft_delete_all_monitored_articles()
            logger.info("全量软删除文章完成 count=%s", deleted_count)
            return deleted_count
        except Exception:
            logger.exception("全量软删除文章失败")
            raise

    def delete_monitored_article(self, article_id: int) -> bool:
        logger.info("开始软删除文章 id=%s", article_id)
        try:
            deleted = delete_monitored_article(article_id)
            logger.info("文章软删除完成 id=%s success=%s", article_id, deleted)
            return deleted
        except Exception:
            logger.exception("文章软删除失败 id=%s", article_id)
            raise

    def rewrite_title(self, title: str) -> str:
        logger.info("开始改写标题")
        try:
            result = TitleRewriter().rewrite(title)
            logger.info("标题改写完成 length=%s", len(result))
            return result
        except Exception:
            logger.exception("标题改写失败")
            raise

    def rewrite_article_full(self, article_id: int, template_key: str) -> dict:
        logger.info("收到完整改写请求 article_id=%s template_key=%s", article_id, template_key)
        try:
            rewrite_workers = self._load_worker_count(REWRITE_WORKER_COUNT_KEY, DEFAULT_REWRITE_WORKER_COUNT)
            image_gen_workers = self._load_worker_count(IMAGE_GEN_WORKER_COUNT_KEY, DEFAULT_IMAGE_GEN_WORKER_COUNT)
            pipeline = ArticleRewritePipeline(
                rewrite_workers=rewrite_workers,
                image_gen_workers=image_gen_workers,
            )
            result = pipeline.run(article_id, template_key)
            logger.info("完整改写完成 article_id=%s", article_id)
            return result
        except Exception:
            logger.exception("完整改写失败 article_id=%s template_key=%s", article_id, template_key)
            raise

    @staticmethod
    def _load_worker_count(key: str, default: int) -> int:
        raw = get_setting(key)
        try:
            count = int(raw) if raw is not None else default
        except ValueError:
            count = default
        return max(1, min(count, MAX_WORKER_COUNT))

    def open_external_url(self, url: str) -> bool:
        try:
            return webbrowser.open(url)
        except Exception:
            return False

    def open_logs_folder(self) -> bool:
        success = _open_folder(LOG_DIR)
        logger.info("打开日志目录 success=%s path=%s", success, LOG_DIR)
        return success

    def open_db_folder(self) -> bool:
        success = _open_folder(DB_DIR)
        logger.info("打开数据库目录 success=%s path=%s", success, DB_DIR)
        return success

    def clear_logs(self) -> bool:
        logger.info("开始清理日志文件 path=%s", LOG_DIR)
        LOG_DIR.mkdir(parents=True, exist_ok=True)
        removed = False

        for file in LOG_DIR.glob("app.log*"):
            if not file.is_file():
                continue
            try:
                os.remove(file)
                removed = True
            except OSError:
                logger.warning("删除日志文件失败 path=%s", file)
                continue

        logger.info("日志文件清理完成 removed=%s", removed)
        return removed

