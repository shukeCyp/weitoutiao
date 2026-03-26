from __future__ import annotations

from concurrent.futures import ThreadPoolExecutor, as_completed
import json
import logging
import re
from dataclasses import dataclass
from html import unescape
from pathlib import Path
from queue import Empty, Queue
from threading import Lock
from typing import Any, Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from playwright.sync_api import TimeoutError as PlaywrightTimeoutError
from playwright.sync_api import BrowserContext, sync_playwright

from .db import MonitoredArticle, get_setting
from .monitoring import (
    AUTOMATION_WORKER_COUNT_KEY,
    DEFAULT_HEADLESS,
    DEFAULT_USER_AGENT,
    DEFAULT_WORKER_COUNT,
    MAX_WORKER_COUNT,
    PAGE_BOOTSTRAP_TIMEOUT_MS,
)
from .title_rewriter import ContentTitleGenerator

logger = logging.getLogger("weitoutiao.article_exporter")

AUTOMATION_HEADLESS_KEY = "automation.headless"
DOWNLOAD_TITLE_PROMPT_KEY = "article_download.title_prompt"
DEFAULT_EXPORT_TIMEOUT_MS = PAGE_BOOTSTRAP_TIMEOUT_MS
DEFAULT_EXPORT_WAIT_MS = 3000
MIN_VALID_CONTENT_LENGTH = 80
MAX_FILENAME_LENGTH = 80


@dataclass(slots=True)
class ExporterSettings:
    headless: bool
    worker_count: int


class OriginalArticleBatchExporter:
    def __init__(self) -> None:
        self.settings = self._load_settings()
        self.title_generator = ContentTitleGenerator()
        self._save_path_lock = Lock()

    def export_articles(
        self,
        articles: list[MonitoredArticle],
        folder: Path,
        *,
        progress_callback: Callable[[dict[str, Any]], None] | None = None,
    ) -> dict[str, Any]:
        folder.mkdir(parents=True, exist_ok=True)

        progress: dict[str, Any] = {
            "succeeded": 0,
            "failed": 0,
            "skipped": 0,
            "errors": [],
            "done": 0,
            "total": len(articles),
            "currentArticle": "",
        }
        progress_lock = Lock()

        def emit_progress() -> None:
            if progress_callback:
                progress_callback(dict(progress))

        if not articles:
            emit_progress()
            return dict(progress)

        article_queue: Queue[MonitoredArticle] = Queue()
        for article in articles:
            article_queue.put(article)

        reserved_paths: set[str] = set()

        def set_current_article(article: MonitoredArticle) -> None:
            preview = (article.content or "").strip().replace("\n", " ")
            current = preview[:28] if preview else f"ID {article.id}"
            with progress_lock:
                progress["currentArticle"] = current
                emit_progress()

        def mark_finished(article: MonitoredArticle, status: str, *, error: str | None = None) -> None:
            with progress_lock:
                progress["done"] += 1
                if status == "success":
                    progress["succeeded"] += 1
                elif status == "skipped":
                    progress["skipped"] += 1
                else:
                    progress["failed"] += 1
                    if error:
                        progress["errors"] = [*progress["errors"], f"ID {article.id}: {error}"][:5]
                emit_progress()

        def worker() -> None:
            with sync_playwright() as playwright:
                browser = playwright.chromium.launch(headless=self.settings.headless)
                context = browser.new_context(user_agent=DEFAULT_USER_AGENT, locale="zh-CN")

                try:
                    while True:
                        try:
                            article = article_queue.get_nowait()
                        except Empty:
                            break

                        set_current_article(article)
                        try:
                            title, content, source_url = self._prepare_article(context, article)
                            if not title or not content:
                                logger.warning("原文导出跳过：缺少标题或正文 article_id=%s", article.id)
                                mark_finished(article, "skipped")
                                continue

                            with self._save_path_lock:
                                save_path = self._build_unique_path(folder, title, article.id, reserved_paths)
                                reserved_paths.add(str(save_path))

                            self._build_docx(
                                title=title,
                                content=content,
                                source_url=source_url,
                                publish_time=article.publish_time.isoformat() if article.publish_time else None,
                                save_path=save_path,
                            )
                            logger.info("原文导出成功 article_id=%s path=%s", article.id, save_path)
                            mark_finished(article, "success")
                        except Exception as exc:  # noqa: BLE001
                            logger.warning("原文导出失败 article_id=%s error=%s", article.id, exc, exc_info=True)
                            mark_finished(article, "failed", error=str(exc))
                        finally:
                            article_queue.task_done()
                finally:
                    context.close()
                    browser.close()

        emit_progress()
        worker_count = max(1, min(self.settings.worker_count, len(articles)))
        with ThreadPoolExecutor(max_workers=worker_count, thread_name_prefix="original-export") as executor:
            futures = [executor.submit(worker) for _ in range(worker_count)]
            for future in as_completed(futures):
                future.result()

        with progress_lock:
            progress["currentArticle"] = ""
            emit_progress()
            return dict(progress)

    def _load_settings(self) -> ExporterSettings:
        headless_raw = get_setting(AUTOMATION_HEADLESS_KEY)
        worker_count_raw = get_setting(AUTOMATION_WORKER_COUNT_KEY)

        headless = DEFAULT_HEADLESS if headless_raw is None else headless_raw.lower() == "true"
        try:
            worker_count = int(worker_count_raw) if worker_count_raw is not None else DEFAULT_WORKER_COUNT
        except ValueError:
            worker_count = DEFAULT_WORKER_COUNT

        worker_count = max(1, min(worker_count, MAX_WORKER_COUNT))
        return ExporterSettings(headless=headless, worker_count=worker_count)

    def _prepare_article(self, context: BrowserContext, article: MonitoredArticle) -> tuple[str, str, str | None]:
        source_url = self._resolve_source_url(article)
        fallback_content = self._extract_fallback_content(article)

        scraped_content = ""
        if source_url:
            try:
                scraped_content = self._scrape_article_content(context, source_url)
            except Exception as exc:  # noqa: BLE001
                logger.warning("Playwright 抓取正文失败 article_id=%s url=%s error=%s", article.id, source_url, exc)

        content = self._choose_best_content(scraped_content, fallback_content)
        if not content:
            raise RuntimeError("未能获取有效正文。")

        prompt_template = get_setting(DOWNLOAD_TITLE_PROMPT_KEY)
        try:
            title = self.title_generator.generate(content, prompt_template)
        except Exception as exc:  # noqa: BLE001
            logger.warning("模型标题生成失败 article_id=%s error=%s，回退到正文首行", article.id, exc)
            title = self._fallback_title(content, article)

        return title.strip(), content.strip(), source_url

    def _resolve_source_url(self, article: MonitoredArticle) -> str | None:
        candidates = [article.display_url]

        raw = self._parse_raw_json(article.raw_json)
        if raw:
            candidates.extend(
                [
                    self._get_nested_text(raw, "share_url"),
                    self._get_nested_text(raw, "share", "share_url"),
                    self._get_nested_text(raw, "share_info", "share_url"),
                    self._get_nested_text(raw, "shareInfo", "shareURL"),
                ]
            )

        for candidate in candidates:
            if isinstance(candidate, str) and candidate.strip().startswith(("http://", "https://")):
                return candidate.strip()
        return None

    def _scrape_article_content(self, context: BrowserContext, url: str) -> str:
        page = context.new_page()
        try:
            logger.info("开始抓取文章正文 url=%s", url)
            page.goto(url, wait_until="domcontentloaded", timeout=DEFAULT_EXPORT_TIMEOUT_MS)
            page.wait_for_timeout(DEFAULT_EXPORT_WAIT_MS)
            try:
                page.wait_for_load_state("networkidle", timeout=5000)
            except PlaywrightTimeoutError:
                logger.info("文章页未进入 networkidle，继续解析 url=%s", url)

            payload = page.evaluate(
                """
                () => {
                  const normalize = (text) => (text || '')
                    .replace(/\\u00a0/g, ' ')
                    .replace(/\\r/g, '')
                    .split('\\n')
                    .map((line) => line.trim())
                    .filter(Boolean)
                    .join('\\n\\n')
                    .trim()

                  const isVisible = (el) => {
                    if (!(el instanceof HTMLElement)) return false
                    const style = window.getComputedStyle(el)
                    const rect = el.getBoundingClientRect()
                    return style.display !== 'none' && style.visibility !== 'hidden' && rect.width > 0 && rect.height > 0
                  }

                  const candidates = []
                  const selectors = [
                    'article',
                    'main',
                    '[class*="article"]',
                    '[class*="content"]',
                    '[class*="detail"]',
                    '[class*="rich"]',
                    '[class*="thread"]',
                    '[data-testid*="content"]',
                  ]

                  for (const selector of selectors) {
                    for (const node of document.querySelectorAll(selector)) {
                      if (!(node instanceof HTMLElement) || !isVisible(node)) continue
                      const text = normalize(node.innerText)
                      if (!text) continue
                      const marker = `${node.id || ''} ${node.className || ''}`.toLowerCase()
                      let score = text.length
                      if (/article|content|detail|rich|thread|main/.test(marker)) score += 300
                      if (/comment|footer|recommend|related|toolbar|action|nav|header|aside/.test(marker)) score -= 400
                      candidates.push({ text, score })
                    }
                  }

                  candidates.sort((a, b) => b.score - a.score)

                  return {
                    best: candidates[0]?.text || '',
                    body: normalize(document.body?.innerText || ''),
                    meta: normalize(document.querySelector('meta[name="description"]')?.content || ''),
                  }
                }
                """
            )

            if not isinstance(payload, dict):
                return ""

            best = self._clean_text(str(payload.get("best") or ""))
            body = self._clean_text(str(payload.get("body") or ""))
            meta = self._clean_text(str(payload.get("meta") or ""))

            if len(best) >= MIN_VALID_CONTENT_LENGTH:
                return best
            if len(body) >= MIN_VALID_CONTENT_LENGTH:
                return body
            return meta
        finally:
            page.close()

    def _extract_fallback_content(self, article: MonitoredArticle) -> str:
        candidates: list[str] = []
        if article.content:
            candidates.append(article.content)

        raw = self._parse_raw_json(article.raw_json)
        if raw:
            candidates.extend(
                [
                    self._get_nested_text(raw, "content"),
                    self._get_nested_text(raw, "rich_content"),
                    self._get_nested_text(raw, "richContent"),
                    self._get_nested_text(raw, "itemCell", "richContentInfo", "richContent"),
                ]
            )

        cleaned_candidates = [self._clean_text(text) for text in candidates if isinstance(text, str) and text.strip()]
        if not cleaned_candidates:
            return ""
        cleaned_candidates.sort(key=len, reverse=True)
        return cleaned_candidates[0]

    def _choose_best_content(self, scraped_content: str, fallback_content: str) -> str:
        scraped = self._clean_text(scraped_content)
        fallback = self._clean_text(fallback_content)

        if scraped and fallback:
            if self._content_matches(scraped, fallback):
                return scraped if len(scraped) >= len(fallback) * 0.7 else fallback
            return scraped if len(scraped) >= max(MIN_VALID_CONTENT_LENGTH, len(fallback) * 1.6) else fallback
        if scraped:
            return scraped
        return fallback

    def _content_matches(self, scraped: str, fallback: str) -> bool:
        fallback_anchor = self._content_anchor(fallback)
        normalized_scraped = re.sub(r"\s+", "", scraped)
        normalized_fallback = re.sub(r"\s+", "", fallback)
        if not fallback_anchor or not normalized_scraped:
            return False
        return fallback_anchor in normalized_scraped or self._content_anchor(scraped) in normalized_fallback

    def _content_anchor(self, content: str) -> str:
        normalized = re.sub(r"\s+", "", content)
        return normalized[:30]

    def _fallback_title(self, content: str, article: MonitoredArticle) -> str:
        lines = [line.strip() for line in content.splitlines() if line.strip()]
        if lines:
            return lines[0][:40]
        if article.content:
            return article.content.strip()[:40]
        return f"文章_{article.id}"

    def _parse_raw_json(self, raw_json: str | None) -> dict[str, Any] | None:
        if not raw_json:
            return None
        try:
            parsed = json.loads(raw_json)
        except json.JSONDecodeError:
            return None
        return parsed if isinstance(parsed, dict) else None

    def _get_nested_text(self, payload: dict[str, Any], *keys: str) -> str | None:
        current: Any = payload
        for key in keys:
            if not isinstance(current, dict):
                return None
            current = current.get(key)
        return current.strip() if isinstance(current, str) and current.strip() else None

    def _clean_text(self, value: str) -> str:
        normalized = unescape(value or "")
        normalized = re.sub(r"<br\s*/?>", "\n", normalized, flags=re.IGNORECASE)
        normalized = re.sub(r"</p\s*>", "\n", normalized, flags=re.IGNORECASE)
        normalized = re.sub(r"<[^>]+>", "", normalized)
        normalized = normalized.replace("\r", "")
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        normalized = "\n".join(line.strip() for line in normalized.splitlines())
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()

    def _build_unique_path(self, folder: Path, title: str, article_id: int, reserved_paths: set[str]) -> Path:
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", title).strip()[:MAX_FILENAME_LENGTH] or f"文章_{article_id}"
        save_path = folder / f"{safe_name}.docx"
        if not save_path.exists() and str(save_path) not in reserved_paths:
            return save_path

        stem = save_path.stem
        suffix = save_path.suffix
        counter = 1
        while save_path.exists() or str(save_path) in reserved_paths:
            save_path = folder / f"{stem}_{counter}{suffix}"
            counter += 1
        return save_path

    @staticmethod
    def _build_docx(
        *,
        title: str,
        content: str,
        source_url: str | None,
        publish_time: str | None,
        save_path: Path,
    ) -> None:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title.strip())
        title_run.bold = True
        title_run.font.size = Pt(16)

        if publish_time:
            meta_para = doc.add_paragraph(f"发布时间：{publish_time}")
            meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in meta_para.runs:
                run.font.size = Pt(10.5)

        if source_url:
            source_para = doc.add_paragraph(f"原文链接：{source_url}")
            source_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in source_para.runs:
                run.font.size = Pt(10.5)

        if publish_time or source_url:
            doc.add_paragraph()

        for paragraph in [line.strip() for line in content.split("\n") if line.strip()]:
            text_para = doc.add_paragraph(paragraph)
            text_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in text_para.runs:
                run.font.size = Pt(12)

        save_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(save_path))


__all__ = ["DOWNLOAD_TITLE_PROMPT_KEY", "OriginalArticleBatchExporter"]
